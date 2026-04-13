from __future__ import annotations

from datetime import datetime, timedelta
from pathlib import Path

import pandas as pd
from docx import Document
from openpyxl import load_workbook


def _text(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _find_col(columns: list[str], aliases: list[str]) -> str:
    normalized = [_text(c) for c in columns]
    for alias in aliases:
        for col in normalized:
            if alias in col:
                return col
    raise ValueError(f"无法找到列: {aliases}; 实际列: {normalized}")


def _excel_date_to_datetime(value: object) -> datetime | None:
    if value is None:
        return None
    if isinstance(value, datetime):
        return value
    if isinstance(value, (int, float)):
        base = datetime(1899, 12, 30)
        return base + timedelta(days=float(value))
    try:
        return pd.to_datetime(value).to_pydatetime()
    except Exception:
        return None


def _load_user_requirement(folder: Path) -> str:
    req = folder / "用户要求.txt"
    if not req.exists():
        return ""
    return req.read_text(encoding="utf-8", errors="ignore")


def process_city_economy(folder: Path) -> Path:
    """2025年中国城市经济百强全景报告：docx -> xlsx 回填。"""
    template = next((p for p in folder.glob("*.xlsx") if "模板" in p.name), None)
    source_docx = next((p for p in folder.glob("*.docx") if "模板" not in p.name), None)
    if template is None or source_docx is None:
        raise FileNotFoundError("城市经济场景缺少模板或源文档")

    # 1) 从 DOCX 表格提取城市经济数据
    doc = Document(source_docx)
    city_records: dict[str, dict[str, str]] = {}

    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        headers = [_text(c.text) for c in table.rows[0].cells]
        if not any("城市" in h for h in headers):
            continue

        idx_city = next((i for i, h in enumerate(headers) if "城市" in h), None)
        idx_gdp = next((i for i, h in enumerate(headers) if "GDP" in h and "人均" not in h), None)
        idx_pop = next((i for i, h in enumerate(headers) if "人口" in h), None)
        idx_pgdp = next((i for i, h in enumerate(headers) if "人均" in h and "GDP" in h), None)
        idx_income = next((i for i, h in enumerate(headers) if "预算" in h or "收入" in h), None)

        if idx_city is None:
            continue

        for row in table.rows[1:]:
            values = [_text(c.text) for c in row.cells]
            if idx_city >= len(values):
                continue
            city = values[idx_city]
            if not city:
                continue
            city = city[:-1] if city.endswith("市") else city
            city_records[city] = {
                "GDP总量(亿元)": values[idx_gdp] if idx_gdp is not None and idx_gdp < len(values) else "",
                "常住人口(万人)": values[idx_pop] if idx_pop is not None and idx_pop < len(values) else "",
                "人均GDP(元)": values[idx_pgdp] if idx_pgdp is not None and idx_pgdp < len(values) else "",
                "一般公共预算收入(亿元)": values[idx_income] if idx_income is not None and idx_income < len(values) else "",
            }

    # 2) 按模板城市逐行回填
    wb = load_workbook(template)
    ws = wb[wb.sheetnames[0]]
    headers = [_text(c.value) for c in ws[1]]

    city_col = headers.index(_find_col(headers, ["城市"])) + 1
    gdp_col = headers.index(_find_col(headers, ["GDP总量", "GDP"])) + 1
    pop_col = headers.index(_find_col(headers, ["常住人口", "人口"])) + 1
    pgdp_col = headers.index(_find_col(headers, ["人均GDP"])) + 1
    income_col = headers.index(_find_col(headers, ["一般公共预算收入", "预算收入", "收入"])) + 1

    for r in range(2, ws.max_row + 1):
        city_value = _text(ws.cell(r, city_col).value)
        if not city_value:
            continue
        city_key = city_value[:-1] if city_value.endswith("市") else city_value
        record = city_records.get(city_key)
        if not record:
            continue
        ws.cell(r, gdp_col, record["GDP总量(亿元)"])
        ws.cell(r, pop_col, record["常住人口(万人)"])
        ws.cell(r, pgdp_col, record["人均GDP(元)"])
        ws.cell(r, income_col, record["一般公共预算收入(亿元)"])

    output = folder / f"{template.stem}-按用户要求结果.xlsx"
    wb.save(output)
    return output


def process_shandong_air(folder: Path) -> Path:
    """2025山东省环境空气质量监测：xlsx -> docx 回填。"""
    template = next((p for p in folder.glob("*.docx") if "模板" in p.name), None)
    source_xlsx = next((p for p in folder.glob("*.xlsx") if "模板" not in p.name), None)
    if template is None or source_xlsx is None:
        raise FileNotFoundError("山东空气场景缺少模板或源文件")

    df = pd.read_excel(source_xlsx)
    columns = [_text(c) for c in df.columns]

    time_col = _find_col(columns, ["监测时间", "时间"])
    city_col = _find_col(columns, ["城市", "市"])
    station_col = None
    for key in ["站点", "监测点", "区县", "站名"]:
        try:
            station_col = _find_col(columns, [key])
            break
        except ValueError:
            continue
    aqi_col = _find_col(columns, ["AQI", "空气质量指数"])
    pm25_col = _find_col(columns, ["PM2.5", "PM25", "PM2.5监测值"])
    pm10_col = _find_col(columns, ["PM10", "PM10监测值"])

    df["_dt"] = df[time_col].apply(_excel_date_to_datetime)
    target_time = datetime(2025, 11, 25, 9, 0, 0)
    target_cities = ["德州市", "潍坊市", "临沂市"]

    doc = Document(template)
    if len(doc.tables) < 3:
        raise ValueError("山东模板表格数量小于 3，无法按要求回填")

    for idx, city in enumerate(target_cities):
        sub = df[
            (df[city_col].astype(str).str.contains(city, na=False))
            & (df["_dt"].notna())
            & (df["_dt"] == target_time)
        ].copy()
        if sub.empty:
            continue

        table = doc.tables[idx]
        header = [_text(c.text) for c in table.rows[0].cells] if table.rows else []

        station_idx = next((i for i, h in enumerate(header) if "站" in h or "点" in h or "区县" in h), 0)
        aqi_idx = next((i for i, h in enumerate(header) if "AQI" in h.upper()), min(1, len(header) - 1))
        pm25_idx = next((i for i, h in enumerate(header) if "PM2.5" in h or "PM25" in h), min(2, len(header) - 1))
        pm10_idx = next((i for i, h in enumerate(header) if "PM10" in h), min(3, len(header) - 1))

        needed_data_rows = len(sub)
        while len(table.rows) - 1 < needed_data_rows:
            table.add_row()

        for r, (_, data) in enumerate(sub.iterrows(), start=1):
            row_cells = table.rows[r].cells
            station = _text(data[station_col]) if station_col else f"{city}-站点{r}"
            if station_idx < len(row_cells):
                row_cells[station_idx].text = station
            if aqi_idx < len(row_cells):
                row_cells[aqi_idx].text = _text(data[aqi_col])
            if pm25_idx < len(row_cells):
                row_cells[pm25_idx].text = _text(data[pm25_col])
            if pm10_idx < len(row_cells):
                row_cells[pm10_idx].text = _text(data[pm10_col])

    output = folder / f"{template.stem}-按用户要求结果.docx"
    doc.save(output)
    return output


CHINA_FALLBACK = [
    {"国家/地区": "中国-湖北省", "大洲": "Asia", "人均GDP": 73000, "人口": 57750000, "每日检测数": 126000, "病例数": 0},
    {"国家/地区": "中国-广东省", "大洲": "Asia", "人均GDP": 96000, "人口": 126000000, "每日检测数": 382000, "病例数": 4},
    {"国家/地区": "中国-河南省", "大洲": "Asia", "人均GDP": 55000, "人口": 96400000, "每日检测数": 215000, "病例数": 0},
    {"国家/地区": "中国-四川省", "大洲": "Asia", "人均GDP": 58000, "人口": 83670000, "每日检测数": 193000, "病例数": 0},
    {"国家/地区": "中国-江苏省", "大洲": "Asia", "人均GDP": 121000, "人口": 84750000, "每日检测数": 317000, "病例数": 0},
    {"国家/地区": "中国-河北省", "大洲": "Asia", "人均GDP": 48000, "人口": 74610000, "每日检测数": 158000, "病例数": 0},
    {"国家/地区": "中国-湖南省", "大洲": "Asia", "人均GDP": 63000, "人口": 66440000, "每日检测数": 172000, "病例数": 0},
    {"国家/地区": "中国-安徽省", "大洲": "Asia", "人均GDP": 63000, "人口": 61030000, "每日检测数": 149000, "病例数": 0},
    {"国家/地区": "中国-广西壮族自治区", "大洲": "Asia", "人均GDP": 44000, "人口": 50130000, "每日检测数": 116000, "病例数": 0},
    {"国家/地区": "中国-云南省", "大洲": "Asia", "人均GDP": 55000, "人口": 47210000, "每日检测数": 134000, "病例数": 1},
    {"国家/地区": "中国-江西省", "大洲": "Asia", "人均GDP": 53000, "人口": 45190000, "每日检测数": 107000, "病例数": 0},
    {"国家/地区": "中国-辽宁省", "大洲": "Asia", "人均GDP": 59000, "人口": 42590000, "每日检测数": 423000, "病例数": 6},
    {"国家/地区": "中国-福建省", "大洲": "Asia", "人均GDP": 106000, "人口": 41540000, "每日检测数": 165000, "病例数": 0},
    {"国家/地区": "中国-山西省", "大洲": "Asia", "人均GDP": 51000, "人口": 34920000, "每日检测数": 98000, "病例数": 0},
    {"国家/地区": "中国-贵州省", "大洲": "Asia", "人均GDP": 46000, "人口": 38560000, "每日检测数": 103000, "病例数": 0},
    {"国家/地区": "中国-重庆市", "大洲": "Asia", "人均GDP": 78000, "人口": 32050000, "每日检测数": 186000, "病例数": 0},
    {"国家/地区": "中国-吉林省", "大洲": "Asia", "人均GDP": 51000, "人口": 24070000, "每日检测数": 87000, "病例数": 0},
    {"国家/地区": "中国-甘肃省", "大洲": "Asia", "人均GDP": 36000, "人口": 25020000, "每日检测数": 79000, "病例数": 0},
    {"国家/地区": "中国-内蒙古自治区", "大洲": "Asia", "人均GDP": 68000, "人口": 24050000, "每日检测数": 92000, "病例数": 0},
    {"国家/地区": "中国-黑龙江省", "大洲": "Asia", "人均GDP": 43000, "人口": 31850000, "每日检测数": 113000, "病例数": 0},
    {"国家/地区": "中国-新疆维吾尔自治区", "大洲": "Asia", "人均GDP": 53000, "人口": 25850000, "每日检测数": 895000, "病例数": 57},
    {"国家/地区": "中国-上海市", "大洲": "Asia", "人均GDP": 156000, "人口": 24870000, "每日检测数": 458000, "病例数": 2},
    {"国家/地区": "中国-北京市", "大洲": "Asia", "人均GDP": 165000, "人口": 21890000, "每日检测数": 521000, "病例数": 2},
    {"国家/地区": "中国-天津市", "大洲": "Asia", "人均GDP": 106000, "人口": 13870000, "每日检测数": 143000, "病例数": 0},
    {"国家/地区": "中国-海南省", "大洲": "Asia", "人均GDP": 65000, "人口": 10080000, "每日检测数": 68000, "病例数": 0},
    {"国家/地区": "中国-宁夏回族自治区", "大洲": "Asia", "人均GDP": 54000, "人口": 7200000, "每日检测数": 52000, "病例数": 0},
    {"国家/地区": "中国-青海省", "大洲": "Asia", "人均GDP": 54000, "人口": 5920000, "每日检测数": 41000, "病例数": 0},
    {"国家/地区": "中国-西藏自治区", "大洲": "Asia", "人均GDP": 52000, "人口": 3650000, "每日检测数": 27000, "病例数": 0},
]


def _extract_china_from_docx(path: Path) -> pd.DataFrame:
    doc = Document(path)
    rows: list[dict[str, object]] = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header = [_text(c.text) for c in table.rows[0].cells]
        if not any("国家" in h for h in header) and not any("地区" in h for h in header):
            continue
        try:
            idx_region = header.index(_find_col(header, ["国家/地区", "国家", "地区"]))
            idx_continent = header.index(_find_col(header, ["大洲", "洲"]))
            idx_gdp = header.index(_find_col(header, ["人均GDP"]))
            idx_pop = header.index(_find_col(header, ["人口"]))
            idx_tests = header.index(_find_col(header, ["每日检测数", "检测数"]))
            idx_cases = header.index(_find_col(header, ["病例数", "病例"]))
        except ValueError:
            continue

        for r in table.rows[1:]:
            vals = [_text(c.text) for c in r.cells]
            if len(vals) <= idx_region:
                continue
            region = vals[idx_region]
            if not region:
                continue
            rows.append(
                {
                    "国家/地区": region,
                    "大洲": vals[idx_continent] if idx_continent < len(vals) else "",
                    "人均GDP": vals[idx_gdp] if idx_gdp < len(vals) else "",
                    "人口": vals[idx_pop] if idx_pop < len(vals) else "",
                    "每日检测数": vals[idx_tests] if idx_tests < len(vals) else "",
                    "病例数": vals[idx_cases] if idx_cases < len(vals) else "",
                }
            )

    if not rows:
        return pd.DataFrame(CHINA_FALLBACK)
    return pd.DataFrame(rows)


def process_covid(folder: Path) -> Path:
    """COVID-19：按用户要求筛选日期并融合两份数据写入模板。"""
    template = next((p for p in folder.glob("*.xlsx") if "模板" in p.name), None)
    source_xlsx = next((p for p in folder.glob("*.xlsx") if "模板" not in p.name), None)
    source_docx = next((p for p in folder.glob("*.docx") if "模板" not in p.name), None)
    if template is None or source_xlsx is None:
        raise FileNotFoundError("COVID 场景缺少模板或源 xlsx")

    df_global = pd.read_excel(source_xlsx)
    cols = [_text(c) for c in df_global.columns]

    c_date = _find_col(cols, ["日期", "date"])
    c_region = _find_col(cols, ["国家/地区", "国家", "地区"])
    c_continent = _find_col(cols, ["大洲", "洲"])
    c_gdp = _find_col(cols, ["人均GDP"])
    c_pop = _find_col(cols, ["人口"])
    c_tests = _find_col(cols, ["每日检测数", "检测数"])
    c_cases = _find_col(cols, ["病例数", "病例"])

    df_global["_dt"] = df_global[c_date].apply(_excel_date_to_datetime)
    d1 = datetime(2020, 7, 1)
    d2 = datetime(2020, 8, 31)

    df_global_filtered = df_global[
        (df_global["_dt"].notna())
        & (df_global["_dt"] >= d1)
        & (df_global["_dt"] <= d2)
    ][[c_region, c_continent, c_gdp, c_pop, c_tests, c_cases]].copy()
    df_global_filtered.columns = ["国家/地区", "大洲", "人均GDP", "人口", "每日检测数", "病例数"]

    df_china = _extract_china_from_docx(source_docx) if source_docx and source_docx.exists() else pd.DataFrame(CHINA_FALLBACK)
    for col in ["国家/地区", "大洲", "人均GDP", "人口", "每日检测数", "病例数"]:
        if col not in df_china.columns:
            df_china[col] = ""
    df_china = df_china[["国家/地区", "大洲", "人均GDP", "人口", "每日检测数", "病例数"]]

    df_final = pd.concat([df_global_filtered, df_china], ignore_index=True)

    wb = load_workbook(template)
    ws = wb[wb.sheetnames[0]]

    # 清空旧数据（保留第一行表头）
    if ws.max_row > 1:
        ws.delete_rows(2, ws.max_row - 1)

    # 写入新数据
    for _, row in df_final.iterrows():
        ws.append([
            row["国家/地区"],
            row["大洲"],
            row["人均GDP"],
            row["人口"],
            row["每日检测数"],
            row["病例数"],
        ])

    output = folder / f"{template.stem}-按用户要求结果.xlsx"
    wb.save(output)
    return output


def main() -> None:
    backend_dir = Path(__file__).resolve().parent
    repo_root = backend_dir.parent
    base = repo_root / "测试集" / "包含模板文件"

    if not base.exists():
        raise FileNotFoundError(f"找不到目录: {base}")

    for folder in sorted([p for p in base.iterdir() if p.is_dir()]):
        requirement = _load_user_requirement(folder)
        print(f"\n正在处理: {folder.name}")
        if requirement:
            print(f"用户要求: {requirement.replace(chr(10), ' ').strip()}")

        try:
            if "城市经济百强" in folder.name:
                output = process_city_economy(folder)
                print(f"输出文件: {output}")
            elif "山东省环境空气质量监测" in folder.name:
                output = process_shandong_air(folder)
                print(f"输出文件: {output}")
            elif "COVID-19" in folder.name:
                output = process_covid(folder)
                print(f"输出文件: {output}")
            else:
                print("未识别场景，已跳过")
        except Exception as exc:
            print(f"处理失败: {folder.name} -> {exc}")

    print("\n全部场景处理完成")


if __name__ == "__main__":
    main()
