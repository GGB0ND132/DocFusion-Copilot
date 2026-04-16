from __future__ import annotations

import zipfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree as ET

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"
W = {"w": W_NS}

ET.register_namespace("w", W_NS)


@dataclass(slots=True)
class WordTableRow:
    """Word 表格中单行的内存表示。    In-memory representation of one DOCX table row."""

    row_index: int
    values: list[str]


@dataclass(slots=True)
class WordTable:
    """Word 文档中单张表格的内存表示。    In-memory representation of one DOCX table."""

    table_index: int
    name: str
    rows: list[WordTableRow]
    context_text: str = ""


@dataclass(slots=True)
class WordDocument:
    """DOCX 文档表格视图的内存表示。    In-memory representation of the table view of a DOCX document."""

    tables: list[WordTable]


@dataclass(slots=True)
class WordCellWrite:
    """Word 表格单元格写入动作。    Planned write operation for one DOCX table cell."""

    table_index: int
    row_index: int
    column_index: int
    value: str | float | int


def _w(tag: str) -> str:
    """为 WordprocessingML 标签补齐命名空间。    Qualify a WordprocessingML tag with its namespace."""

    return f"{{{W_NS}}}{tag}"


def _parse_table_rows(table_el: ET.Element) -> list[WordTableRow]:
    """解析单张表格的所有行。    Parse all rows from one table element."""

    rows: list[WordTableRow] = []
    for row_index, row_el in enumerate(table_el.findall("w:tr", W), start=1):
        values: list[str] = []
        logical_col = 0
        for cell_el in row_el.findall("w:tc", W):
            tc_pr = cell_el.find("w:tcPr", W)
            # Check for vertical merge continuation — skip these cells' content
            v_merge = tc_pr.find("w:vMerge", W) if tc_pr is not None else None
            is_v_merge_continue = False
            if v_merge is not None:
                val = v_merge.get(_w("val"), "")
                if val != "restart":
                    is_v_merge_continue = True

            grid_span = 1
            if tc_pr is not None:
                gs_el = tc_pr.find("w:gridSpan", W)
                if gs_el is not None:
                    try:
                        grid_span = int(gs_el.get(_w("val"), "1"))
                    except (ValueError, TypeError):
                        grid_span = 1

            text = "" if is_v_merge_continue else _text_from_element(cell_el).strip()

            while len(values) < logical_col:
                values.append("")
            values.append(text)
            for _ in range(grid_span - 1):
                values.append("")
            logical_col += grid_span

        rows.append(WordTableRow(row_index=row_index, values=values))
    return rows


def load_docx_tables(path: str | Path) -> WordDocument:
    """读取 DOCX 中的表格结构，正确处理合并单元格，并收集每张表前的段落文字。
    Read DOCX table structures into memory, handling merged cells correctly,
    and capture preceding paragraph text for each table."""

    docx_path = Path(path)
    with zipfile.ZipFile(docx_path, "r") as archive:
        root = ET.fromstring(archive.read("word/document.xml"))

    body = root.find(_w("body"))
    if body is None:
        return WordDocument(tables=[])

    tables: list[WordTable] = []
    pending_paragraphs: list[str] = []
    table_counter = 0
    tbl_tag = _w("tbl")
    p_tag = _w("p")

    for child in body:
        if child.tag == p_tag:
            text = _text_from_element(child).strip()
            if text:
                pending_paragraphs.append(text)
        elif child.tag == tbl_tag:
            table_counter += 1
            context_text = "\n".join(pending_paragraphs[-6:])
            pending_paragraphs.clear()
            rows = _parse_table_rows(child)
            tables.append(WordTable(
                table_index=table_counter,
                name=f"table_{table_counter}",
                rows=rows,
                context_text=context_text,
            ))

    return WordDocument(tables=tables)


def apply_docx_updates(template_path: str | Path, output_path: str | Path, updates: list[WordCellWrite]) -> None:
    """将单元格更新写回 DOCX 表格（基于 python-docx）。
    Apply cell updates back into a DOCX table document using python-docx."""
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn

    source_file = Path(template_path)
    destination_file = Path(output_path)
    destination_file.parent.mkdir(parents=True, exist_ok=True)

    doc = DocxDocument(str(source_file))
    tables = doc.tables

    grouped_updates: dict[int, list[WordCellWrite]] = {}
    for update in updates:
        grouped_updates.setdefault(update.table_index, []).append(update)

    for table_index, table_updates in grouped_updates.items():
        if table_index < 1 or table_index > len(tables):
            raise ValueError(f"Table index out of range: {table_index}")

        table = tables[table_index - 1]

        # Determine how many rows / columns we need
        max_row = max(u.row_index for u in table_updates)
        max_col = max(u.column_index for u in table_updates)

        # Add missing rows by cloning the last row (preserving column structure)
        while len(table.rows) < max_row:
            last_row_el = table.rows[-1]._tr
            new_row_el = deepcopy(last_row_el)
            # Clear text and vertical merge continuations in cloned cells
            for tc in new_row_el.findall(qn("w:tc")):
                for p in tc.findall(qn("w:p")):
                    for r in p.findall(qn("w:r")):
                        p.remove(r)
                tc_pr = tc.find(qn("w:tcPr"))
                if tc_pr is not None:
                    v_merge = tc_pr.find(qn("w:vMerge"))
                    if v_merge is not None:
                        tc_pr.remove(v_merge)
            table._tbl.append(new_row_el)

        for update in sorted(table_updates, key=lambda u: (u.row_index, u.column_index)):
            row = table.rows[update.row_index - 1]
            # Add missing cells if column doesn't exist
            while len(row.cells) < update.column_index:
                last_cell_el = row.cells[-1]._tc
                new_cell_el = deepcopy(last_cell_el)
                for p in new_cell_el.findall(qn("w:p")):
                    for r in p.findall(qn("w:r")):
                        p.remove(r)
                row._tr.append(new_cell_el)

            cell = row.cells[update.column_index - 1]
            # Preserve first paragraph's formatting, replace text
            if cell.paragraphs:
                para = cell.paragraphs[0]
                # Clear existing runs
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = str(update.value)
                else:
                    para.add_run(str(update.value))
                # Remove extra paragraphs
                for extra_para in cell.paragraphs[1:]:
                    extra_para._element.getparent().remove(extra_para._element)
            else:
                cell.text = str(update.value)

    doc.save(str(destination_file))


def replace_text_in_docx_document(
    source_path: str | Path,
    output_path: str | Path,
    replacements: list[tuple[str, str]],
) -> int:
    """对 DOCX 文档执行简单文本替换。    Apply simple text replacements to a DOCX document."""

    template_file = Path(source_path)
    destination_file = Path(output_path)
    destination_file.parent.mkdir(parents=True, exist_ok=True)

    total_changes = 0
    with zipfile.ZipFile(template_file, "r") as source_archive:
        root = ET.fromstring(source_archive.read("word/document.xml"))
        for paragraph_el in root.findall(".//w:p", W):
            text = _text_from_element(paragraph_el)
            if not text:
                continue
            updated_text, change_count = _apply_replacements(text, replacements)
            if change_count <= 0:
                continue
            _set_paragraph_text(paragraph_el, updated_text)
            total_changes += change_count

        document_payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        with zipfile.ZipFile(destination_file, "w", compression=zipfile.ZIP_DEFLATED) as destination_archive:
            for file_info in source_archive.infolist():
                payload = (
                    document_payload
                    if file_info.filename == "word/document.xml"
                    else source_archive.read(file_info.filename)
                )
                destination_archive.writestr(file_info, payload)
    return total_changes


def _text_from_element(element: ET.Element) -> str:
    """提取 XML 元素中的可见文本。    Extract visible text from one XML element."""

    return "".join(node.text or "" for node in element.findall(".//w:t", W))


def _apply_replacements(text: str, replacements: list[tuple[str, str]]) -> tuple[str, int]:
    """顺序应用文本替换并统计变更次数。    Apply replacements sequentially and count changes."""

    updated_text = text
    total_changes = 0
    for old_text, new_text in replacements:
        change_count = updated_text.count(old_text)
        if change_count <= 0:
            continue
        updated_text = updated_text.replace(old_text, new_text)
        total_changes += change_count
    return updated_text, total_changes


def _get_or_create_table_row(table_el: ET.Element, row_index: int, column_count: int) -> ET.Element:
    """返回指定索引的表格行，不存在时自动补齐。    Return a table row by index, creating missing rows when needed."""

    rows = table_el.findall("w:tr", W)
    while len(rows) < row_index:
        if rows:
            new_row = deepcopy(rows[-1])
            for cell_el in new_row.findall("w:tc", W):
                _set_cell_text(cell_el, "")
                # Clear vMerge from cloned cells to prevent structural corruption
                tc_pr = cell_el.find("w:tcPr", W)
                if tc_pr is not None:
                    v_merge = tc_pr.find("w:vMerge", W)
                    if v_merge is not None:
                        tc_pr.remove(v_merge)
        else:
            new_row = _new_empty_row(column_count)
        table_el.append(new_row)
        rows = table_el.findall("w:tr", W)
    return rows[row_index - 1]


def _get_or_create_table_cell(row_el: ET.Element, column_index: int) -> ET.Element:
    """返回指定列的表格单元格，不存在时自动补齐。    Return a table cell by column index, creating missing cells when needed."""

    cells = row_el.findall("w:tc", W)
    while len(cells) < column_index:
        if cells:
            new_cell = deepcopy(cells[-1])
            _set_cell_text(new_cell, "")
        else:
            new_cell = _new_empty_cell()
        row_el.append(new_cell)
        cells = row_el.findall("w:tc", W)
    return cells[column_index - 1]


def _new_empty_row(column_count: int) -> ET.Element:
    """创建一行空白表格行。    Create a new empty DOCX table row."""

    row_el = ET.Element(_w("tr"))
    for _ in range(max(column_count, 1)):
        row_el.append(_new_empty_cell())
    return row_el


def _new_empty_cell() -> ET.Element:
    """创建一个空白表格单元格。    Create a new empty DOCX table cell."""

    cell_el = ET.Element(_w("tc"))
    paragraph_el = ET.SubElement(cell_el, _w("p"))
    run_el = ET.SubElement(paragraph_el, _w("r"))
    ET.SubElement(run_el, _w("t"))
    return cell_el


def _set_cell_text(cell_el: ET.Element, value: str) -> None:
    """写入单元格文本并尽量保留样式。    Write cell text while preserving existing styling when possible."""

    tc_pr = deepcopy(cell_el.find("w:tcPr", W))
    first_paragraph = cell_el.find("w:p", W)
    paragraph_pr = deepcopy(first_paragraph.find("w:pPr", W)) if first_paragraph is not None else None
    first_run = first_paragraph.find("w:r", W) if first_paragraph is not None else None
    run_pr = deepcopy(first_run.find("w:rPr", W)) if first_run is not None else None

    for child in list(cell_el):
        cell_el.remove(child)

    if tc_pr is not None:
        cell_el.append(tc_pr)

    paragraph_el = ET.SubElement(cell_el, _w("p"))
    if paragraph_pr is not None:
        paragraph_el.append(paragraph_pr)
    run_el = ET.SubElement(paragraph_el, _w("r"))
    if run_pr is not None:
        run_el.append(run_pr)
    text_el = ET.SubElement(run_el, _w("t"))
    if value != value.strip():
        text_el.set(f"{{{XML_NS}}}space", "preserve")
    text_el.text = value


def _set_paragraph_text(paragraph_el: ET.Element, value: str) -> None:
    """重写段落文本并尽量保留段落样式。    Rewrite paragraph text while preserving paragraph styling when possible."""

    paragraph_pr = paragraph_el.find("w:pPr", W)
    run_pr = None
    for run_el in paragraph_el.findall("w:r", W):
        run_pr = deepcopy(run_el.find("w:rPr", W))
        if run_pr is not None:
            break

    for child in list(paragraph_el):
        if child.tag != _w("pPr"):
            paragraph_el.remove(child)

    if paragraph_pr is None:
        paragraph_pr = ET.Element(_w("pPr"))
        paragraph_el.insert(0, paragraph_pr)

    run_el = ET.SubElement(paragraph_el, _w("r"))
    if run_pr is not None:
        run_el.append(run_pr)
    text_el = ET.SubElement(run_el, _w("t"))
    if value != value.strip():
        text_el.set(f"{{{XML_NS}}}space", "preserve")
    text_el.text = value
