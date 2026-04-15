"""端到端基准测试脚本 — 模拟竞赛评分流程。

Usage:
    python scripts/run_benchmark.py [--base-url http://127.0.0.1:8000] [--testset-dir ../测试集]

流程:
    Phase A  一次性上传全部源文档 → 等待解析完成 → 记录总解析时间
    Phase B  逐个场景提交模板回填 → 等待完成 → 下载结果 → 与正确回填文件逐单元格对比
    Phase C  输出汇总报告 (控制台 + JSON)
"""
from __future__ import annotations

import argparse
import json
import sys
import tempfile
import time
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path

import requests
from docx import Document as DocxDocument
from openpyxl import load_workbook

# ---------------------------------------------------------------------------
# 数据结构
# ---------------------------------------------------------------------------

@dataclass
class CellComparison:
    location: str          # e.g. "Sheet1!B3" or "Table1[2,3]"
    expected: str
    actual: str
    match: bool


@dataclass
class ScenarioResult:
    name: str
    template_file: str
    fill_time_seconds: float  = 0.0
    total_blanks: int         = 0
    correct_count: int        = 0
    accuracy: float           = 0.0
    passed: bool              = False
    error: str | None         = None
    cell_details: list[CellComparison] = field(default_factory=list)


@dataclass
class BenchmarkReport:
    timestamp: str
    base_url: str
    parse_time_seconds: float               = 0.0
    document_count: int                     = 0
    document_set_id: str                    = ""
    parse_results: list[dict]               = field(default_factory=list)
    scenarios: list[ScenarioResult]         = field(default_factory=list)
    avg_accuracy: float                     = 0.0
    avg_fill_time: float                    = 0.0


# ---------------------------------------------------------------------------
# 辅助常量 / 工具
# ---------------------------------------------------------------------------

SCENARIO_CONFIGS = [
    {
        "name": "2025年中国城市经济百强全景报告",
        "folder": "2025年中国城市经济百强全景报告",
        "template_glob": "*模板*.xlsx",
        "reference_name": "正确回填文件.xlsx",
        "format": "xlsx",
    },
    {
        "name": "2025山东省环境空气质量监测数据信息",
        "folder": "2025山东省环境空气质量监测数据信息",
        "template_glob": "*模板*.docx",
        "reference_name": "正确回填文件.docx",
        "format": "docx",
    },
    {
        "name": "COVID-19数据集",
        "folder": "COVID-19数据集",
        "template_glob": "*模板*.xlsx",
        "reference_name": "正确回填文件.xlsx",
        "format": "xlsx",
    },
]

POLL_INTERVAL = 2          # 秒
POLL_TIMEOUT  = 900        # 秒 (单个任务最大等待)
_NUM_TOLERANCE = 0.01      # 数值相对容差 1% (可由 --tolerance 覆盖)


def _get_tolerance() -> float:
    return _NUM_TOLERANCE


def _normalize(raw: object) -> str:
    """将单元格值规范化为可比较的字符串。"""
    if raw is None:
        return ""
    s = str(raw).strip()
    # 去除千分位逗号
    s = s.replace(",", "").replace("，", "")
    # 去掉尾部百分号 (在数值对比里再特殊处理)
    return s


def _try_float(s: str) -> float | None:
    """尝试将字符串解析为浮点数，支持百分号。"""
    s = s.strip()
    if not s:
        return None
    pct = False
    if s.endswith("%"):
        s = s[:-1].strip()
        pct = True
    s = s.replace(",", "").replace("，", "")
    try:
        v = float(s)
        return v / 100.0 if pct else v
    except ValueError:
        return None


def _values_match(expected_raw: str, actual_raw: str) -> bool:
    """判断两个值是否匹配。"""
    e = _normalize(expected_raw)
    a = _normalize(actual_raw)
    if e == a:
        return True
    # 尝试数值对比
    ef = _try_float(e)
    af = _try_float(a)
    if ef is not None and af is not None:
        denom = max(abs(ef), 1.0)
        return abs(ef - af) / denom < _get_tolerance()
    return False


def _is_blank(value: object) -> bool:
    """判断单元格是否为空白。"""
    if value is None:
        return True
    s = str(value).strip()
    return s == "" or s == "None"


# ---------------------------------------------------------------------------
# Phase A : 文档上传与解析
# ---------------------------------------------------------------------------

def collect_source_documents(testset_dir: Path) -> list[Path]:
    """从测试集中收集全部源文档（含场景文件夹中的非模板、非用户要求文件）。"""
    files: list[Path] = []
    seen: set[str] = set()

    # 1. 公用源文件（Excel / md / txt / word）
    subdirs = ["Excel", "md", "txt", "word"]
    for sub in subdirs:
        d = testset_dir / sub
        if not d.is_dir():
            print(f"  [WARN] 子目录不存在: {d}")
            continue
        for f in sorted(d.iterdir()):
            if f.is_file() and not f.name.startswith("~$"):
                files.append(f)
                seen.add(f.name)

    # 2. 场景专属源文件（包含模板文件/<场景>/ 中非模板、非用户要求的文件）
    templates_dir = testset_dir / "包含模板文件"
    if templates_dir.is_dir():
        for cfg in SCENARIO_CONFIGS:
            scenario_dir = templates_dir / cfg["folder"]
            if not scenario_dir.is_dir():
                continue
            for f in sorted(scenario_dir.iterdir()):
                if not f.is_file() or f.name.startswith("~$"):
                    continue
                # 跳过模板文件、用户要求、正确回填文件、README
                name_lower = f.name.lower()
                if "模板" in f.name:
                    continue
                if f.name == "用户要求.txt":
                    continue
                if f.name.startswith("正确回填"):
                    continue
                if name_lower in ("readme.txt",):
                    continue
                # 避免重名文件重复上传
                if f.name in seen:
                    continue
                files.append(f)
                seen.add(f.name)
                print(f"  [INFO] 发现场景专属源文件: {f.relative_to(testset_dir)}")

    return files


def upload_batch(base_url: str, files: list[Path]) -> tuple[str, list[dict]]:
    """POST /api/v1/documents/upload-batch — 返回 (document_set_id, items)。"""
    url = f"{base_url}/api/v1/documents/upload-batch"
    multipart = []
    for f in files:
        multipart.append(("files", (f.name, f.read_bytes())))
    resp = requests.post(url, files=multipart, timeout=120)
    resp.raise_for_status()
    data = resp.json()
    return data["document_set_id"], data["items"]


def poll_task(base_url: str, task_id: str, timeout: float = POLL_TIMEOUT) -> dict:
    """轮询直到任务完成，返回最终 task 状态。"""
    url = f"{base_url}/api/v1/tasks/{task_id}"
    t0 = time.time()
    while True:
        resp = requests.get(url, timeout=30)
        resp.raise_for_status()
        task = resp.json()
        status = task.get("status", "")
        if status in ("succeeded", "completed", "success", "failed"):
            return task
        if time.time() - t0 > timeout:
            raise TimeoutError(f"Task {task_id} 超时 ({timeout}s)")
        time.sleep(POLL_INTERVAL)


def phase_a(base_url: str, testset_dir: Path) -> tuple[str, float, list[dict]]:
    """Phase A: 上传 + 解析。返回 (document_set_id, parse_elapsed, parse_results)。"""
    print("\n" + "=" * 60)
    print("Phase A: 文档上传与解析")
    print("=" * 60)

    files = collect_source_documents(testset_dir)
    print(f"  收集到 {len(files)} 个源文档:")
    for f in files:
        print(f"    - {f.name}")

    print(f"\n  正在上传 {len(files)} 个文件…")
    t_start = time.time()
    doc_set_id, items = upload_batch(base_url, files)
    print(f"  上传完成。document_set_id = {doc_set_id}")
    print(f"  等待全部解析任务完成…")

    parse_results: list[dict] = []
    for item in items:
        tid = item["task_id"]
        fname = item["document"]["file_name"]
        try:
            task = poll_task(base_url, tid)
            parse_results.append({
                "file_name": fname,
                "task_id": tid,
                "status": task["status"],
                "message": task.get("message", ""),
            })
        except Exception as exc:
            parse_results.append({
                "file_name": fname,
                "task_id": tid,
                "status": "error",
                "message": str(exc),
            })

    parse_elapsed = time.time() - t_start
    succeeded = sum(1 for r in parse_results if r["status"] in ("succeeded", "completed", "success"))
    failed = len(parse_results) - succeeded
    print(f"\n  解析完成: {succeeded} 成功, {failed} 失败, 总耗时 {parse_elapsed:.1f}s")
    for r in parse_results:
        icon = "✓" if r["status"] in ("succeeded", "completed", "success") else "✗"
        print(f"    [{icon}] {r['file_name']} — {r['status']}")

    return doc_set_id, parse_elapsed, parse_results


# ---------------------------------------------------------------------------
# Phase B : 模板回填 + 准确率评估
# ---------------------------------------------------------------------------

def submit_fill(base_url: str, template_path: Path, doc_set_id: str,
                user_requirement: str) -> str:
    """POST /api/v1/templates/fill — 返回 task_id。"""
    url = f"{base_url}/api/v1/templates/fill"
    resp = requests.post(
        url,
        files={"template_file": (template_path.name, template_path.read_bytes())},
        data={
            "document_set_id": doc_set_id,
            "fill_mode": "canonical",
            "auto_match": "true",
            "user_requirement": user_requirement,
        },
        timeout=30,
    )
    resp.raise_for_status()
    return resp.json()["task_id"]


def download_result(base_url: str, task_id: str, dest: Path) -> Path:
    """GET /api/v1/templates/result/{task_id} — 下载到 dest 文件。"""
    url = f"{base_url}/api/v1/templates/result/{task_id}"
    resp = requests.get(url, timeout=60)
    resp.raise_for_status()

    # 从 Content-Disposition 推断文件名，否则保持原名
    cd = resp.headers.get("content-disposition", "")
    if "filename=" in cd:
        fname = cd.split("filename=")[-1].strip().strip('"')
        dest = dest.parent / fname

    dest.write_bytes(resp.content)
    return dest


# -- XLSX 对比 --

def compare_xlsx(template_path: Path, reference_path: Path,
                 actual_path: Path) -> list[CellComparison]:
    """XLSX 逐单元格对比: 模板空 + 参考有值 → blank → 对比实际。"""
    wb_tpl = load_workbook(template_path, data_only=True)
    wb_ref = load_workbook(reference_path, data_only=True)
    wb_act = load_workbook(actual_path, data_only=True)

    results: list[CellComparison] = []
    for sheet_name in wb_ref.sheetnames:
        ws_ref = wb_ref[sheet_name]
        ws_tpl = wb_tpl[sheet_name] if sheet_name in wb_tpl.sheetnames else None
        ws_act = wb_act[sheet_name] if sheet_name in wb_act.sheetnames else None

        for row in ws_ref.iter_rows():
            for cell in row:
                ref_val = cell.value
                if _is_blank(ref_val):
                    continue
                # 检查模板对应位置是否为空 (blank)
                tpl_val = ws_tpl[cell.coordinate].value if ws_tpl else None
                if not _is_blank(tpl_val):
                    continue  # 模板已有值，不算 blank

                # 这是一个待填 blank
                act_val = ws_act[cell.coordinate].value if ws_act else None
                expected_str = _normalize(str(ref_val))
                actual_str = _normalize(str(act_val)) if act_val is not None else ""
                match = _values_match(str(ref_val), str(act_val) if act_val is not None else "")
                results.append(CellComparison(
                    location=f"{sheet_name}!{cell.coordinate}",
                    expected=expected_str,
                    actual=actual_str,
                    match=match,
                ))

    wb_tpl.close()
    wb_ref.close()
    wb_act.close()
    return results


# -- DOCX 对比 --

def _docx_table_cells(doc: DocxDocument) -> list[tuple[int, int, int, str]]:
    """提取 DOCX 中所有表格单元格，去重合并单元格。
    返回 [(table_idx, row_idx, col_idx, text), ...]
    """
    seen: set[tuple[int, int, int]] = set()
    cells: list[tuple[int, int, int, str]] = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                key = (t_idx, r_idx, c_idx)
                # python-docx 在合并区域返回相同的 cell 对象
                # 通过 (tc element id) 去重
                tc_id = id(cell._tc)
                # 但同一合并单元格在不同 (r,c) 位置对应同一 _tc
                # 只保留首次出现的 (r,c) 位置
                if key in seen:
                    continue
                seen.add(key)
                cells.append((t_idx, r_idx, c_idx, cell.text.strip()))
    return cells


def compare_docx(template_path: Path, reference_path: Path,
                 actual_path: Path) -> list[CellComparison]:
    """DOCX 逐表格单元格对比。"""
    doc_tpl = DocxDocument(str(template_path))
    doc_ref = DocxDocument(str(reference_path))
    doc_act = DocxDocument(str(actual_path))

    cells_tpl = {(t, r, c): text for t, r, c, text in _docx_table_cells(doc_tpl)}
    cells_ref = {(t, r, c): text for t, r, c, text in _docx_table_cells(doc_ref)}
    cells_act = {(t, r, c): text for t, r, c, text in _docx_table_cells(doc_act)}

    results: list[CellComparison] = []
    for key, ref_text in cells_ref.items():
        if not ref_text:
            continue
        tpl_text = cells_tpl.get(key, "")
        if tpl_text:
            continue  # 模板已有值，不算 blank

        act_text = cells_act.get(key, "")
        t_idx, r_idx, c_idx = key
        match = _values_match(ref_text, act_text)
        results.append(CellComparison(
            location=f"Table{t_idx + 1}[{r_idx},{c_idx}]",
            expected=_normalize(ref_text),
            actual=_normalize(act_text),
            match=match,
        ))

    return results


def phase_b_scenario(base_url: str, doc_set_id: str, testset_dir: Path,
                     cfg: dict) -> ScenarioResult:
    """执行单个场景的回填 + 评估。"""
    scenario_dir = testset_dir / "包含模板文件" / cfg["folder"]
    result = ScenarioResult(name=cfg["name"], template_file="")

    # 1) 定位文件
    template_path = next(scenario_dir.glob(cfg["template_glob"]), None)
    reference_path = scenario_dir / cfg["reference_name"]
    requirement_path = scenario_dir / "用户要求.txt"

    if template_path is None:
        result.error = f"模板文件未找到: {cfg['template_glob']}"
        return result
    if not reference_path.exists():
        result.error = f"参考答案未找到: {reference_path}"
        return result

    result.template_file = template_path.name
    user_req = requirement_path.read_text(encoding="utf-8", errors="ignore").strip() if requirement_path.exists() else ""

    print(f"\n  场景: {cfg['name']}")
    print(f"    模板: {template_path.name}")
    print(f"    需求: {user_req[:80]}{'…' if len(user_req) > 80 else ''}")

    # 2) 提交回填
    t_start = time.time()
    try:
        task_id = submit_fill(base_url, template_path, doc_set_id, user_req)
    except Exception as exc:
        result.error = f"提交回填失败: {exc}"
        return result

    print(f"    task_id = {task_id}")

    # 3) 轮询至完成
    try:
        task = poll_task(base_url, task_id, timeout=POLL_TIMEOUT)
    except Exception as exc:
        result.fill_time_seconds = time.time() - t_start
        result.error = f"回填任务失败: {exc}"
        return result

    fill_elapsed = time.time() - t_start
    result.fill_time_seconds = fill_elapsed

    if task["status"] not in ("succeeded", "completed", "success"):
        result.error = f"回填任务状态: {task['status']} — {task.get('error', task.get('message', ''))}"
        return result

    print(f"    回填完成: {fill_elapsed:.1f}s")

    # 4) 下载结果
    with tempfile.TemporaryDirectory() as tmpdir:
        dest = Path(tmpdir) / f"result.{cfg['format']}"
        try:
            actual_path = download_result(base_url, task_id, dest)
        except Exception as exc:
            result.error = f"下载结果失败: {exc}"
            return result

        # 5) 逐单元格对比
        try:
            if cfg["format"] == "xlsx":
                comparisons = compare_xlsx(template_path, reference_path, actual_path)
            else:
                comparisons = compare_docx(template_path, reference_path, actual_path)
        except Exception as exc:
            result.error = f"对比失败: {exc}"
            return result

    result.cell_details = comparisons
    result.total_blanks = len(comparisons)
    result.correct_count = sum(1 for c in comparisons if c.match)
    result.accuracy = (result.correct_count / result.total_blanks) if result.total_blanks > 0 else 0.0
    result.passed = result.accuracy >= 0.80 and result.fill_time_seconds <= 90.0

    print(f"    Blanks: {result.total_blanks}, 正确: {result.correct_count}, "
          f"准确率: {result.accuracy:.1%}, "
          f"{'✓ PASS' if result.passed else '✗ FAIL'}")

    # 打印不匹配的 cells (最多20个)
    mismatches = [c for c in comparisons if not c.match]
    if mismatches:
        print(f"    不匹配 ({len(mismatches)}):")
        for c in mismatches[:20]:
            print(f"      {c.location}: 期望={c.expected!r}  实际={c.actual!r}")
        if len(mismatches) > 20:
            print(f"      ... 还有 {len(mismatches) - 20} 个")

    return result


# ---------------------------------------------------------------------------
# Phase C : 汇总报告
# ---------------------------------------------------------------------------

def phase_c(report: BenchmarkReport, report_dir: Path) -> Path:
    """生成控制台汇总 + JSON 报告。"""
    print("\n" + "=" * 60)
    print("Phase C: 汇总报告")
    print("=" * 60)

    # -- 解析阶段摘要 --
    print(f"\n  文档解析: {report.document_count} 个文件, 总耗时 {report.parse_time_seconds:.1f}s")

    # -- 回填汇总表格 --
    divider = "-" * 90
    print(f"\n  {divider}")
    print(f"  {'场景':<30s} {'Blanks':>7s} {'正确':>6s} {'准确率':>8s} {'时间(s)':>8s} {'达标':>6s}")
    print(f"  {divider}")

    for s in report.scenarios:
        if s.error:
            print(f"  {s.name:<30s} {'ERROR':>7s} {'':<6s} {'':<8s} {s.fill_time_seconds:>8.1f} {'✗':>6s}")
            print(f"    错误: {s.error}")
        else:
            status = "✓" if s.passed else "✗"
            print(f"  {s.name:<30s} {s.total_blanks:>7d} {s.correct_count:>6d} "
                  f"{s.accuracy:>7.1%} {s.fill_time_seconds:>8.1f} {status:>6s}")

    print(f"  {divider}")

    valid = [s for s in report.scenarios if s.error is None]
    if valid:
        report.avg_accuracy = sum(s.accuracy for s in valid) / len(valid)
        report.avg_fill_time = sum(s.fill_time_seconds for s in valid) / len(valid)
        print(f"  {'平均':<30s} {'':<7s} {'':<6s} {report.avg_accuracy:>7.1%} {report.avg_fill_time:>8.1f}")
    else:
        print(f"  所有场景均失败，无法计算平均值。")

    print(f"  {divider}\n")

    # -- 保存 JSON --
    report_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    report_path = report_dir / f"benchmark_{ts}.json"

    # 序列化
    payload = asdict(report)
    report_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str),
                           encoding="utf-8")
    print(f"  JSON 报告已保存: {report_path}")
    return report_path


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------

def main() -> None:
    global _NUM_TOLERANCE

    parser = argparse.ArgumentParser(description="DocFusion 端到端基准测试")
    parser.add_argument("--base-url", default="http://127.0.0.1:8000",
                        help="后端 API 地址 (默认 http://127.0.0.1:8000)")
    parser.add_argument("--testset-dir", default=None,
                        help="测试集根目录 (默认 <workspace>/测试集)")
    parser.add_argument("--report-dir", default=None,
                        help="报告输出目录 (默认 storage/benchmark_reports)")
    parser.add_argument("--tolerance", type=float, default=_NUM_TOLERANCE,
                        help=f"数值对比相对容差 (默认 {_NUM_TOLERANCE})")
    args = parser.parse_args()

    _NUM_TOLERANCE = args.tolerance

    # 推断路径
    backend_dir = Path(__file__).resolve().parent.parent
    workspace_dir = backend_dir.parent

    testset_dir = Path(args.testset_dir) if args.testset_dir else workspace_dir / "测试集"
    report_dir = Path(args.report_dir) if args.report_dir else backend_dir / "storage" / "benchmark_reports"

    if not testset_dir.is_dir():
        print(f"[ERROR] 测试集目录不存在: {testset_dir}")
        sys.exit(1)

    base_url = args.base_url.rstrip("/")
    print(f"DocFusion 端到端基准测试")
    print(f"  后端地址: {base_url}")
    print(f"  测试集:   {testset_dir}")
    print(f"  报告目录: {report_dir}")

    report = BenchmarkReport(
        timestamp=datetime.now(timezone.utc).isoformat(),
        base_url=base_url,
    )

    # Phase A
    doc_set_id, parse_elapsed, parse_results = phase_a(base_url, testset_dir)
    report.document_set_id = doc_set_id
    report.parse_time_seconds = parse_elapsed
    report.document_count = len(parse_results)
    report.parse_results = parse_results

    # Phase B
    print("\n" + "=" * 60)
    print("Phase B: 模板回填与准确率评估")
    print("=" * 60)

    for cfg in SCENARIO_CONFIGS:
        sr = phase_b_scenario(base_url, doc_set_id, testset_dir, cfg)
        report.scenarios.append(sr)

    # Phase C
    phase_c(report, report_dir)

    # 退出码: 全部通过则 0
    all_passed = all(s.passed for s in report.scenarios)
    sys.exit(0 if all_passed else 1)


if __name__ == "__main__":
    main()
