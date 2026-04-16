"""快速测试脚本 — 只跑山东 + COVID 两个场景。

Usage:
    cd backend
    python scripts/run_test.py [--base-url http://127.0.0.1:8000]

输出: 准确率、回填时间、不匹配明细、debug 日志位置。
"""
from __future__ import annotations

import argparse
import glob
import sys
import tempfile
import time
from dataclasses import dataclass, field
from pathlib import Path

import requests
from docx import Document as DocxDocument
from openpyxl import load_workbook

# ---------------------------------------------------------------------------
# 配置
# ---------------------------------------------------------------------------

SCENARIOS = [
    {
        "name": "山东",
        "folder": "2025山东省环境空气质量监测数据信息",
        "template_glob": "*模板*.docx",
        "reference_name": "正确回填文件.docx",
        "format": "docx",
        "sources": [
            "山东省环境空气质量监测数据信息202512171921_0.xlsx",
        ],
    },
    {
        "name": "COVID",
        "folder": "COVID-19数据集",
        "template_glob": "*模板*.xlsx",
        "reference_name": "正确回填文件.xlsx",
        "format": "xlsx",
        "sources": [
            "COVID-19全球数据集（节选）.xlsx",
            "中国COVID-19新冠疫情情况.docx",
        ],
    },
]

POLL_INTERVAL = 2
POLL_TIMEOUT = 180
NUM_TOLERANCE = 0.01
TARGET_ACCURACY = 0.80
TARGET_TIME = 90.0

# ---------------------------------------------------------------------------
# 数据结构
# ---------------------------------------------------------------------------

@dataclass
class CellComparison:
    location: str
    expected: str
    actual: str
    match: bool


@dataclass
class ScenarioResult:
    name: str
    fill_time: float = 0.0
    total_blanks: int = 0
    correct: int = 0
    accuracy: float = 0.0
    passed: bool = False
    error: str | None = None
    mismatches: list[CellComparison] = field(default_factory=list)
    debug_file: str | None = None


# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------

def _normalize(raw: object) -> str:
    if raw is None:
        return ""
    s = str(raw).strip()
    s = s.replace(",", "").replace("，", "")
    return s


def _try_float(s: str) -> float | None:
    s = s.strip()
    if not s:
        return None
    pct = s.endswith("%")
    if pct:
        s = s[:-1].strip()
    s = s.replace(",", "").replace("，", "")
    try:
        v = float(s)
        return v / 100 if pct else v
    except ValueError:
        return None


def _values_match(expected_raw: str, actual_raw: str) -> bool:
    e, a = _normalize(expected_raw), _normalize(actual_raw)
    if e == a:
        return True
    ef, af = _try_float(e), _try_float(a)
    if ef is not None and af is not None:
        denom = max(abs(ef), 1.0)
        return abs(ef - af) / denom < NUM_TOLERANCE
    return False


def _is_blank(value: object) -> bool:
    if value is None:
        return True
    s = str(value).strip()
    return s == "" or s == "None"


# ---------------------------------------------------------------------------
# API helpers
# ---------------------------------------------------------------------------

def upload_files(base_url: str, files: list[Path]) -> tuple[str, list[dict]]:
    url = f"{base_url}/api/v1/documents/upload-batch"
    multipart = [("files", (f.name, f.read_bytes())) for f in files]
    resp = requests.post(url, files=multipart, timeout=120)
    resp.raise_for_status()
    data = resp.json()
    return data["document_set_id"], data["items"]


def poll_task(base_url: str, task_id: str) -> dict:
    url = f"{base_url}/api/v1/tasks/{task_id}"
    t0 = time.time()
    while True:
        resp = requests.get(url, timeout=30)
        resp.raise_for_status()
        task = resp.json()
        status = task.get("status", "")
        if status in ("succeeded", "completed", "success", "failed"):
            return task
        if time.time() - t0 > POLL_TIMEOUT:
            raise TimeoutError(f"Task {task_id} 超时 ({POLL_TIMEOUT}s)")
        time.sleep(POLL_INTERVAL)


def submit_fill(base_url: str, template_path: Path, doc_set_id: str,
                user_requirement: str) -> str:
    url = f"{base_url}/api/v1/templates/fill"
    resp = requests.post(
        url,
        files={"template_file": (template_path.name, template_path.read_bytes())},
        data={
            "document_set_id": doc_set_id,
            "fill_mode": "canonical",
            "auto_match": "true",
            "user_requirement": user_requirement,
        },
        timeout=30,
    )
    resp.raise_for_status()
    return resp.json()["task_id"]


def download_result(base_url: str, task_id: str, dest: Path) -> Path:
    url = f"{base_url}/api/v1/templates/result/{task_id}"
    resp = requests.get(url, timeout=60)
    resp.raise_for_status()
    cd = resp.headers.get("content-disposition", "")
    if "filename=" in cd:
        fname = cd.split("filename=")[-1].strip().strip('"')
        dest = dest.parent / fname
    dest.write_bytes(resp.content)
    return dest


# ---------------------------------------------------------------------------
# 对比
# ---------------------------------------------------------------------------

def _match_rows_greedy(ref_rows: list[dict[str, str]],
                       act_rows: list[dict[str, str]],
                       ) -> list[tuple[dict[str, str], dict[str, str]]]:
    """Greedy row matching: for each ref row, find the best-matching actual row."""
    used: set[int] = set()
    pairs: list[tuple[dict[str, str], dict[str, str]]] = []
    for ref_row in ref_rows:
        best_idx, best_score = -1, 0
        for a_idx, act_row in enumerate(act_rows):
            if a_idx in used:
                continue
            score = sum(1 for c in ref_row
                        if _values_match(ref_row[c], act_row.get(c, "")))
            if score > best_score:
                best_score = score
                best_idx = a_idx
        matched = act_rows[best_idx] if best_idx >= 0 else {}
        if best_idx >= 0:
            used.add(best_idx)
        pairs.append((ref_row, matched))
    return pairs


def compare_xlsx(template_path: Path, reference_path: Path,
                 actual_path: Path) -> list[CellComparison]:
    """Row-order insensitive xlsx comparison."""
    wb_tpl = load_workbook(template_path, data_only=True)
    wb_ref = load_workbook(reference_path, data_only=True)
    wb_act = load_workbook(actual_path, data_only=True)

    results: list[CellComparison] = []
    for sheet_name in wb_ref.sheetnames:
        ws_ref = wb_ref[sheet_name]
        ws_tpl = wb_tpl[sheet_name] if sheet_name in wb_tpl.sheetnames else None
        ws_act = wb_act[sheet_name] if sheet_name in wb_act.sheetnames else None

        # Collect reference data rows: {row_num: {col_letter: (norm_val, coord)}}
        ref_rows_map: dict[int, dict[str, tuple[str, str]]] = {}
        for row in ws_ref.iter_rows():
            for cell in row:
                ref_val = cell.value
                if _is_blank(ref_val):
                    continue
                tpl_val = ws_tpl[cell.coordinate].value if ws_tpl else None
                if not _is_blank(tpl_val):
                    continue
                r = cell.row
                c = cell.column_letter
                if r not in ref_rows_map:
                    ref_rows_map[r] = {}
                ref_rows_map[r][c] = (_normalize(str(ref_val)), cell.coordinate)

        # Collect actual data rows: {row_num: {col_letter: norm_val}}
        act_rows_map: dict[int, dict[str, str]] = {}
        if ws_act:
            for row in ws_act.iter_rows():
                for cell in row:
                    act_val = cell.value
                    if act_val is None:
                        continue
                    r = cell.row
                    c = cell.column_letter
                    if r not in act_rows_map:
                        act_rows_map[r] = {}
                    act_rows_map[r][c] = _normalize(str(act_val))

        # Build matching-compatible row dicts
        ref_row_list = []  # [{col_letter: norm_val}] — for matching
        ref_meta_list = []  # [{col_letter: (norm_val, coord)}] — for reporting
        for rn in sorted(ref_rows_map):
            ref_meta_list.append(ref_rows_map[rn])
            ref_row_list.append({c: v for c, (v, _) in ref_rows_map[rn].items()})

        act_row_list = [act_rows_map[rn] for rn in sorted(act_rows_map)]

        # Greedy row matching
        pairs = _match_rows_greedy(ref_row_list, act_row_list)
        for (ref_row, matched_act), ref_meta in zip(pairs, ref_meta_list):
            for col, (ref_val, coord) in ref_meta.items():
                act_val = matched_act.get(col, "")
                match = _values_match(ref_val, act_val)
                results.append(CellComparison(
                    location=f"{sheet_name}!{coord}",
                    expected=ref_val,
                    actual=act_val,
                    match=match,
                ))

    wb_tpl.close()
    wb_ref.close()
    wb_act.close()
    return results


def compare_docx(template_path: Path, reference_path: Path,
                 actual_path: Path) -> list[CellComparison]:
    """Row-order insensitive docx comparison."""
    doc_tpl = DocxDocument(str(template_path))
    doc_ref = DocxDocument(str(reference_path))
    doc_act = DocxDocument(str(actual_path))

    results: list[CellComparison] = []
    num_tables = len(doc_ref.tables)
    for t_idx in range(num_tables):
        tbl_tpl = doc_tpl.tables[t_idx] if t_idx < len(doc_tpl.tables) else None
        tbl_ref = doc_ref.tables[t_idx]
        tbl_act = doc_act.tables[t_idx] if t_idx < len(doc_act.tables) else None
        if not tbl_tpl or not tbl_act:
            continue

        # Build reference data rows (skip header row 0)
        ref_rows: list[dict[str, str]] = []  # [{col_idx_str: norm_val}]
        ref_meta: list[dict[str, tuple[str, int]]] = []  # [{col_str: (val, r_idx)}]
        seen_ref: set[tuple[int, int]] = set()  # handle merged cells
        for r_idx in range(1, len(tbl_ref.rows)):
            row_vals: dict[str, str] = {}
            row_meta: dict[str, tuple[str, int]] = {}
            for c_idx in range(len(tbl_ref.rows[r_idx].cells)):
                key = (r_idx, c_idx)
                if key in seen_ref:
                    continue
                seen_ref.add(key)
                tpl_val = ""
                if r_idx < len(tbl_tpl.rows) and c_idx < len(tbl_tpl.rows[r_idx].cells):
                    tpl_val = tbl_tpl.rows[r_idx].cells[c_idx].text.strip()
                ref_val = tbl_ref.rows[r_idx].cells[c_idx].text.strip()
                if tpl_val or not ref_val:
                    continue
                cs = str(c_idx)
                row_vals[cs] = _normalize(ref_val)
                row_meta[cs] = (_normalize(ref_val), r_idx)
            if row_vals:
                ref_rows.append(row_vals)
                ref_meta.append(row_meta)

        # Build actual data rows
        act_rows: list[dict[str, str]] = []
        seen_act: set[tuple[int, int]] = set()
        for r_idx in range(1, len(tbl_act.rows)):
            row_vals = {}
            for c_idx in range(len(tbl_act.rows[r_idx].cells)):
                key = (r_idx, c_idx)
                if key in seen_act:
                    continue
                seen_act.add(key)
                act_val = tbl_act.rows[r_idx].cells[c_idx].text.strip()
                row_vals[str(c_idx)] = _normalize(act_val) if act_val else ""
            act_rows.append(row_vals)

        # Greedy row matching
        pairs = _match_rows_greedy(ref_rows, act_rows)
        for (ref_row, matched_act), meta in zip(pairs, ref_meta):
            for cs, (ref_val, orig_r) in meta.items():
                act_val = matched_act.get(cs, "")
                match = _values_match(ref_val, act_val)
                results.append(CellComparison(
                    location=f"Table{t_idx + 1}[{orig_r},{cs}]",
                    expected=ref_val,
                    actual=act_val,
                    match=match,
                ))

    return results


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------

def find_latest_debug(outputs_dir: Path, template_stem: str) -> str | None:
    """找到最新的 debug 文件。"""
    pattern = str(outputs_dir / f"debug_*{template_stem}*.txt")
    files = sorted(glob.glob(pattern), key=lambda f: Path(f).stat().st_mtime, reverse=True)
    return files[0] if files else None


def run_scenario(base_url: str, testset_dir: Path, outputs_dir: Path,
                 cfg: dict, doc_set_id: str) -> ScenarioResult:
    scenario_dir = testset_dir / "包含模板文件" / cfg["folder"]
    result = ScenarioResult(name=cfg["name"])

    template_path = next(scenario_dir.glob(cfg["template_glob"]), None)
    reference_path = scenario_dir / cfg["reference_name"]
    requirement_path = scenario_dir / "用户要求.txt"

    if not template_path:
        result.error = f"模板不存在: {cfg['template_glob']}"
        return result
    if not reference_path.exists():
        result.error = f"参考答案不存在"
        return result

    user_req = requirement_path.read_text("utf-8").strip() if requirement_path.exists() else ""

    print(f"\n{'='*50}")
    print(f"场景: {cfg['name']}")
    print(f"模板: {template_path.name}")
    print(f"需求: {user_req[:100]}")
    print(f"{'='*50}")

    # 提交回填
    t0 = time.time()
    try:
        task_id = submit_fill(base_url, template_path, doc_set_id, user_req)
    except Exception as exc:
        result.error = f"提交失败: {exc}"
        return result
    print(f"  task_id = {task_id}")

    # 等待完成
    try:
        task = poll_task(base_url, task_id)
    except Exception as exc:
        result.fill_time = time.time() - t0
        result.error = f"任务失败: {exc}"
        return result

    result.fill_time = time.time() - t0

    if task["status"] not in ("succeeded", "completed", "success"):
        result.error = f"状态: {task['status']} — {task.get('error', task.get('message', ''))}"
        return result

    print(f"  回填完成: {result.fill_time:.1f}s")

    # 查找 debug 文件
    tpl_stem = template_path.stem
    # 找 task_id 前缀匹配的 debug 文件
    debug_file = find_latest_debug(outputs_dir, "")
    # 更精确: 找包含 task_id 短前缀的
    task_short = task_id[:12]  # task_id 的前12字符
    for f in sorted(outputs_dir.glob(f"debug_*task_{task_short}*.txt"),
                    key=lambda p: p.stat().st_mtime, reverse=True):
        debug_file = str(f)
        break
    else:
        # 退而求其次：找最新的包含模板名的 debug 文件
        debug_file = find_latest_debug(outputs_dir, tpl_stem.replace("task_", "").split("_", 1)[-1] if "task_" in tpl_stem else tpl_stem)

    if debug_file:
        result.debug_file = debug_file
        print(f"  debug文件: {Path(debug_file).name}")

    # 下载结果 & 对比
    with tempfile.TemporaryDirectory() as tmpdir:
        dest = Path(tmpdir) / f"result.{cfg['format']}"
        try:
            actual_path = download_result(base_url, task_id, dest)
        except Exception as exc:
            result.error = f"下载失败: {exc}"
            return result

        try:
            if cfg["format"] == "xlsx":
                comparisons = compare_xlsx(template_path, reference_path, actual_path)
            else:
                comparisons = compare_docx(template_path, reference_path, actual_path)
        except Exception as exc:
            result.error = f"对比失败: {exc}"
            return result

    result.total_blanks = len(comparisons)
    result.correct = sum(1 for c in comparisons if c.match)
    result.accuracy = (result.correct / result.total_blanks) if result.total_blanks else 0.0
    result.passed = result.accuracy >= TARGET_ACCURACY and result.fill_time <= TARGET_TIME
    result.mismatches = [c for c in comparisons if not c.match]

    # 输出结果
    status = "✓ PASS" if result.passed else "✗ FAIL"
    print(f"\n  结果: {status}")
    print(f"  Blanks: {result.total_blanks}, 正确: {result.correct}, 准确率: {result.accuracy:.1%}")
    print(f"  时间: {result.fill_time:.1f}s (限制 {TARGET_TIME:.0f}s)")

    if result.mismatches:
        print(f"\n  不匹配 ({len(result.mismatches)}):")
        for c in result.mismatches[:30]:
            print(f"    {c.location}: 期望={c.expected!r}  实际={c.actual!r}")
        if len(result.mismatches) > 30:
            print(f"    ... 还有 {len(result.mismatches) - 30} 个")

    return result


def main():
    parser = argparse.ArgumentParser(description="快速测试: 山东 + COVID")
    parser.add_argument("--base-url", default="http://127.0.0.1:8000")
    parser.add_argument("--testset-dir", default=None)
    parser.add_argument("--skip-upload", action="store_true",
                        help="跳过上传，使用已有的 document_set_id")
    parser.add_argument("--doc-set-id", default=None,
                        help="已有的 document_set_id (配合 --skip-upload)")
    args = parser.parse_args()

    backend_dir = Path(__file__).resolve().parent.parent
    workspace_dir = backend_dir.parent
    testset_dir = Path(args.testset_dir) if args.testset_dir else workspace_dir / "测试集"
    outputs_dir = backend_dir / "storage" / "outputs"
    base_url = args.base_url.rstrip("/")

    print(f"快速测试: 山东 + COVID")
    print(f"  后端: {base_url}")
    print(f"  测试集: {testset_dir}")

    # Phase A: 上传源文档
    doc_set_id = args.doc_set_id
    if not args.skip_upload or not doc_set_id:
        print(f"\n--- 上传源文档 ---")
        all_sources: list[Path] = []
        seen_names: set[str] = set()
        for cfg in SCENARIOS:
            scenario_dir = testset_dir / "包含模板文件" / cfg["folder"]
            for src_name in cfg["sources"]:
                src_path = scenario_dir / src_name
                if src_path.exists() and src_name not in seen_names:
                    all_sources.append(src_path)
                    seen_names.add(src_name)
                    print(f"  + {src_name}")

        t0 = time.time()
        doc_set_id, items = upload_files(base_url, all_sources)
        print(f"  上传完成 (doc_set_id={doc_set_id})")

        # 等待解析
        for item in items:
            try:
                poll_task(base_url, item["task_id"])
                print(f"  ✓ {item['document']['file_name']}")
            except Exception as exc:
                print(f"  ✗ {item['document']['file_name']}: {exc}")

        print(f"  解析耗时: {time.time() - t0:.1f}s")
    else:
        print(f"\n  跳过上传，使用 doc_set_id={doc_set_id}")

    # Phase B: 回填测试
    results: list[ScenarioResult] = []
    for cfg in SCENARIOS:
        sr = run_scenario(base_url, testset_dir, outputs_dir, cfg, doc_set_id)
        results.append(sr)

    # 汇总
    print(f"\n{'='*60}")
    print(f"{'场景':<12s} {'Blanks':>7s} {'正确':>6s} {'准确率':>8s} {'时间(s)':>8s} {'达标':>6s}")
    print(f"{'-'*60}")
    for sr in results:
        if sr.error:
            print(f"{sr.name:<12s} {'ERROR':>7s} {'':<6s} {'':<8s} {sr.fill_time:>8.1f} {'✗':>6s}")
            print(f"  错误: {sr.error}")
        else:
            icon = "✓" if sr.passed else "✗"
            print(f"{sr.name:<12s} {sr.total_blanks:>7d} {sr.correct:>6d} "
                  f"{sr.accuracy:>7.1%} {sr.fill_time:>8.1f} {icon:>6s}")
    print(f"{'-'*60}")

    valid = [s for s in results if s.error is None]
    if valid:
        avg_acc = sum(s.accuracy for s in valid) / len(valid)
        avg_time = sum(s.fill_time for s in valid) / len(valid)
        print(f"{'平均':<12s} {'':<7s} {'':<6s} {avg_acc:>7.1%} {avg_time:>8.1f}")

    all_passed = all(s.passed for s in results)
    print(f"\n总结: {'全部达标 ✓' if all_passed else '未达标 ✗'}")

    # 显示 debug 文件位置
    for sr in results:
        if sr.debug_file:
            print(f"  {sr.name} debug: {sr.debug_file}")

    sys.exit(0 if all_passed else 1)


if __name__ == "__main__":
    main()
