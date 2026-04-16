# -*- coding: utf-8 -*-
"""
DocFusion Copilot  -  Markdown -> DOCX converter
Reads the project markdown document and generates a formatted Word document.
Competition format: A4, margins 2cm+1cm binding, KaiTi body, single spacing.

All Chinese literals use unicode escapes to avoid encoding issues.
"""
import os
import re
import textwrap
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── paths ──────────────────────────────────────────────────────
BASE = Path(__file__).resolve().parent
MD_FILE = BASE / "\u9879\u76ee\u6587\u6863\u8d44\u6599" / "\u9879\u76ee\u8be6\u7ec6\u65b9\u6848.md"
OUT_FILE = BASE / "\u9879\u76ee\u6587\u6863\u8d44\u6599" / "\u9879\u76ee\u8be6\u7ec6\u65b9\u6848.docx"

# ── font constants (unicode escaped) ───────────────────────────
FONT_BODY   = "\u6977\u4f53"        # KaiTi
FONT_HD     = "\u9ed1\u4f53"        # SimHei
FONT_TITLE  = "\u534e\u6587\u4e2d\u5b8b"  # STZhongsong
FONT_EN     = "Times New Roman"
FONT_CODE   = "Consolas"

# heading sizes
H_SIZES = {1: Pt(16), 2: Pt(14), 3: Pt(13), 4: Pt(12)}
H_SPACE = {
    1: (Pt(24), Pt(12)),
    2: (Pt(18), Pt(8)),
    3: (Pt(12), Pt(6)),
    4: (Pt(8),  Pt(4)),
}


# ── low-level helpers ──────────────────────────────────────────

def _set_cjk(run, font_name):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def mk(para, text, *, font=FONT_BODY, size=Pt(12),
       bold=False, italic=False, color=None, sup=False):
    """Create a styled run inside *para*."""
    r = para.add_run(text)
    r.font.name = font
    _set_cjk(r, font)
    r.font.size = size
    r.font.bold = bold
    r.font.italic = italic
    r.font.superscript = sup
    if color:
        r.font.color.rgb = color
    return r


def body_fmt(para, indent=True):
    pf = para.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after  = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent:
        pf.first_line_indent = Cm(0.74)


def heading_fmt(para, level):
    pf = para.paragraph_format
    sb, sa = H_SPACE.get(level, (Pt(8), Pt(4)))
    pf.space_before = sb
    pf.space_after  = sa
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if level == 1:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Token-level paragraph renderer ────────────────────────────

_TOKEN_RE = re.compile(
    r"(<sup>\[\d+\]</sup>)"   # superscript citation
    r"|(\*\*[^*]+\*\*)"       # bold
    r"|(`[^`]+`)"             # inline code
)


def add_rich_para(doc, text, *, indent=True, font=FONT_BODY, size=Pt(12)):
    """Add a paragraph, rendering **bold**, `code`, and <sup>[n]</sup>."""
    p = doc.add_paragraph()
    body_fmt(p, indent=indent)
    for tok in _TOKEN_RE.split(text):
        if not tok:
            continue
        m_sup = re.match(r"<sup>\[(\d+)\]</sup>", tok)
        if m_sup:
            mk(p, f"[{m_sup.group(1)}]", size=Pt(8), sup=True)
            continue
        if tok.startswith("**") and tok.endswith("**"):
            mk(p, tok[2:-2], bold=True, font=font, size=size)
            continue
        if tok.startswith("`") and tok.endswith("`"):
            mk(p, tok[1:-1], font=FONT_CODE, size=Pt(11))
            continue
        mk(p, tok, font=font, size=size)
    return p


# ── Markdown table renderer ───────────────────────────────────

def add_md_table(doc, header_line, rows):
    """Render a pipe-delimited markdown table as a Word table."""
    headers = [c.strip() for c in header_line.strip().strip("|").split("|")]
    ncols = len(headers)
    tbl = doc.add_table(rows=1, cols=ncols)
    tbl.style = "Table Grid"
    # header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        mk(cell.paragraphs[0], h, font=FONT_HD, size=Pt(10), bold=True)
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # shade header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E3F2FD"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # data rows
    for row_text in rows:
        vals = [c.strip() for c in row_text.strip().strip("|").split("|")]
        row = tbl.add_row()
        for i in range(min(ncols, len(vals))):
            cell = row.cells[i]
            cell.text = ""
            mk(cell.paragraphs[0], vals[i], size=Pt(10))
    # column widths (equal)
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = Cm(14.0 / ncols)
    doc.add_paragraph()  # spacing after table


# ── code block renderer ───────────────────────────────────────

def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.left_indent = Cm(0.5)
        mk(p, line, font=FONT_CODE, size=Pt(8))


# ── figure placeholder ────────────────────────────────────────

def add_figure_caption(doc, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(12)
    # remove **
    clean = caption_text.replace("**", "").strip()
    mk(p, clean, size=Pt(10), bold=True)


# ── main converter ─────────────────────────────────────────────

def convert(md_text, doc):
    """Walk markdown lines and emit Word elements."""
    lines = md_text.split("\n")
    n = len(lines)
    i = 0
    in_mermaid = False
    in_code = False
    code_buf = []
    table_header = None
    table_rows = []
    after_mermaid = False  # flag to catch caption after mermaid block

    def flush_table():
        nonlocal table_header, table_rows
        if table_header:
            add_md_table(doc, table_header, table_rows)
        table_header = None
        table_rows = []

    while i < n:
        line = lines[i]
        raw = line.rstrip()

        # ── mermaid blocks (skip content, render caption after) ──
        if raw.strip().startswith("```mermaid"):
            flush_table()
            in_mermaid = True
            i += 1
            continue
        if in_mermaid:
            if raw.strip() == "```":
                in_mermaid = False
                after_mermaid = True
                # add placeholder
                p = doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after  = Pt(4)
                mk(p, "[\u6b64\u5904\u63d2\u5165\u56fe\u7247]",
                   size=Pt(10), italic=True, color=RGBColor(0x99, 0x99, 0x99))
            i += 1
            continue

        # ── code blocks ──
        if raw.strip().startswith("```") and not in_code:
            flush_table()
            in_code = True
            code_buf = []
            i += 1
            continue
        if in_code:
            if raw.strip() == "```":
                in_code = False
                add_code_block(doc, code_buf)
                code_buf = []
            else:
                code_buf.append(raw)
            i += 1
            continue

        # ── figure caption line (bold, starts with \u56fe) ──
        stripped = raw.strip()
        if stripped.startswith("**\u56fe ") or stripped.startswith("**\u56fe\u00a0") or stripped.startswith("**\u8868 "):
            flush_table()
            add_figure_caption(doc, stripped)
            after_mermaid = False
            i += 1
            continue

        # ── horizontal rule ──
        if stripped == "---":
            flush_table()
            i += 1
            continue

        # ── headings ──
        hm = re.match(r"^(#{1,4})\s+(.+)$", raw)
        if hm:
            flush_table()
            level = len(hm.group(1))
            title = hm.group(2).strip()

            # page break before ## headings (level 2 = chapter)
            if level == 2:
                doc.add_page_break()

            p = doc.add_paragraph()
            heading_fmt(p, level)
            hfont = FONT_HD if level <= 3 else FONT_BODY
            mk(p, title, font=hfont, size=H_SIZES.get(level, Pt(12)), bold=True)
            i += 1
            continue

        # ── markdown table ──
        if "|" in stripped and stripped.startswith("|"):
            # separator line?
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                i += 1
                continue
            if table_header is None:
                table_header = stripped
            else:
                table_rows.append(stripped)
            i += 1
            continue
        else:
            flush_table()

        # ── empty line ──
        if not stripped:
            i += 1
            continue

        # ── skip TOC links ──
        if stripped.startswith("- [") and "](#" in stripped:
            i += 1
            continue

        # ── bold sub-heading like **\uff08\uff09...** ──
        bh = re.match(r"^\*\*[\uff08\(][\d\u4e00-\u9fff]+[\uff09\)].+\*\*$", stripped)
        if bh:
            p = doc.add_paragraph()
            heading_fmt(p, 4)
            mk(p, stripped[2:-2], font=FONT_BODY, size=Pt(12), bold=True)
            i += 1
            continue

        # ── [placeholder] lines ──
        if stripped.startswith("[\u5f85\u8865\u5145") or stripped.startswith("[\u56e2\u961f"):
            p = doc.add_paragraph()
            body_fmt(p)
            mk(p, stripped, italic=True, color=RGBColor(0xFF, 0x66, 0x00))
            i += 1
            continue

        # ── reference entries [N] ... ──
        ref_m = re.match(r"^\[(\d+)\]\s*(.+)", stripped)
        if ref_m:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(1)
            pf.space_after  = Pt(1)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.first_line_indent = Cm(-0.74)
            pf.left_indent = Cm(0.74)
            mk(p, f"[{ref_m.group(1)}] {ref_m.group(2)}",
               font=FONT_EN, size=Pt(9))
            _set_cjk(p.runs[0], FONT_BODY)
            i += 1
            continue

        # ── normal paragraph ──
        add_rich_para(doc, stripped)
        i += 1

    flush_table()


# ── document assembly ──────────────────────────────────────────

def build():
    md_text = MD_FILE.read_text(encoding="utf-8")
    doc = Document()

    # ── page setup ──
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(3)   # 2 + 1 binding
    sec.right_margin  = Cm(2)

    # header
    hp = sec.header.paragraphs[0]
    hp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run(
        "\u3010A23\u3011"
        "\u57fa\u4e8e\u5927\u8bed\u8a00\u6a21\u578b\u7684\u6587\u6863\u7406\u89e3"
        "\u4e0e\u591a\u6e90\u6570\u636e\u878d\u5408\u7cfb\u7edf"
        " \u2014 "
        "\u9879\u76ee\u8be6\u7ec6\u65b9\u6848"
    )
    r.font.name = FONT_BODY
    _set_cjk(r, FONT_BODY)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # footer page number
    fp = sec.footer.paragraphs[0]
    fp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for xml_s in [
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>',
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>',
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>',
    ]:
        r = fp.add_run()
        r._element.append(parse_xml(xml_s))
        r.font.size = Pt(10)

    # ── cover page ──
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mk(p,
       "\u57fa\u4e8e\u5927\u8bed\u8a00\u6a21\u578b\u7684"
       "\u6587\u6863\u7406\u89e3\u4e0e\u591a\u6e90\u6570\u636e\u878d\u5408\u7cfb\u7edf",
       font=FONT_TITLE, size=Pt(22), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    mk(p,
       "DocFusion Copilot \u2014 \u6587\u6863\u667a\u878d\u52a9\u624b",
       font=FONT_HD, size=Pt(15), color=RGBColor(0x4A, 0x90, 0xD9))

    for _ in range(4):
        doc.add_paragraph()
    for line_text in [
        "\u8d5b\u9898\uff1a\u3010A23\u3011"
        "\u57fa\u4e8e\u5927\u8bed\u8a00\u6a21\u578b\u7684\u6587\u6863\u7406\u89e3"
        "\u4e0e\u591a\u6e90\u6570\u636e\u878d\u5408\u7cfb\u7edf"
        "\u3010\u91d1\u9675\u79d1\u6280\u5b66\u9662\u3011",
        "\u7b2c\u5341\u4e03\u5c4a\u4e2d\u56fd\u5927\u5b66\u751f"
        "\u670d\u52a1\u5916\u5305\u521b\u65b0\u521b\u4e1a\u5927\u8d5b",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mk(p, line_text)

    doc.add_page_break()

    # ── skip front-matter in markdown, start from ## ──
    # find first "## " line
    start_idx = 0
    md_lines = md_text.split("\n")
    for idx, ln in enumerate(md_lines):
        if ln.startswith("## ") and "\u76ee\u5f55" not in ln:
            start_idx = idx
            break

    # also skip TOC section (## \u76ee\u5f55 ... ## \u7b2c\u4e00\u7ae0)
    body_text = "\n".join(md_lines[start_idx:])

    convert(body_text, doc)

    doc.save(str(OUT_FILE))
    print(f"\u6587\u6863\u5df2\u4fdd\u5b58\u81f3: {OUT_FILE}")


if __name__ == "__main__":
    build()
